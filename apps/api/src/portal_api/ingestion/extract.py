"""Extração de texto por formato (ADR 0014).

A página é levada até o fim da cadeia porque é ela que aparece na citação. Um
formato sem paginação confiável (``.docx``, ``.md``) devolve página 0, e a
citação sai sem localização — melhor sem do que com uma inventada.

A allowlist é curta de propósito: formato que o portal não sabe ler vira o
estado ``unsupported`` na tela, não uma exceção no log.
"""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass

logger = logging.getLogger(__name__)

MIME_PDF = "application/pdf"
MIME_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
MIME_TEXT = "text/plain"
MIME_MARKDOWN = "text/markdown"
MIME_CSV = "text/csv"

#: O que o upload aceita e a ingestão sabe ler. As duas coisas são a mesma
#: lista de propósito: aceitar no upload o que a ingestão não lê seria prometer
#: uma indexação que nunca acontece.
SUPPORTED_MIME_TYPES: frozenset[str] = frozenset(
    {MIME_PDF, MIME_DOCX, MIME_TEXT, MIME_MARKDOWN, MIME_CSV}
)


class UnsupportedDocument(Exception):
    """Formato fora da allowlist — estado, não falha."""


class ExtractionFailed(Exception):
    """O formato é suportado, mas este arquivo não pôde ser lido."""


@dataclass(frozen=True)
class ExtractedPage:
    """Texto de uma página. ``number`` é 1-based; 0 = formato sem paginação."""

    number: int
    text: str


def extract(data: bytes, mime_type: str | None) -> list[ExtractedPage]:
    """Texto do arquivo, página a página, sem páginas vazias."""
    normalized = (mime_type or "").split(";")[0].strip().lower()
    if normalized not in SUPPORTED_MIME_TYPES:
        raise UnsupportedDocument(normalized or "desconhecido")

    if normalized == MIME_PDF:
        pages = _extract_pdf(data)
    elif normalized == MIME_DOCX:
        pages = _extract_docx(data)
    else:
        pages = _extract_plain(data)

    return [page for page in pages if page.text.strip()]


def _extract_plain(data: bytes) -> list[ExtractedPage]:
    # `replace`: um byte inválido no meio de um relatório não deve custar o
    # documento inteiro — o que dá para ler continua indexável.
    return [ExtractedPage(number=0, text=data.decode("utf-8", errors="replace"))]


def _extract_pdf(data: bytes) -> list[ExtractedPage]:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
        pages = []
        for index, page in enumerate(reader.pages, start=1):
            # Uma página ilegível (só imagem, fonte exótica) não derruba as
            # outras: o documento é indexado pelo que dá para ler.
            try:
                pages.append(ExtractedPage(number=index, text=page.extract_text() or ""))
            except Exception:
                logger.warning("document.page_unreadable", extra={"page": index})
        return pages
    except Exception as exc:  # PDF corrompido ou cifrado
        raise ExtractionFailed("PDF ilegível") from exc


def _extract_docx(data: bytes) -> list[ExtractedPage]:
    import docx

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionFailed("DOCX ilegível") from exc
    # Sem paginação: a quebra de página do .docx é decidida na renderização, e
    # atribuir um número aqui seria inventar a localização da citação.
    text = "\n\n".join(p.text for p in document.paragraphs if p.text.strip())
    return [ExtractedPage(number=0, text=text)]
